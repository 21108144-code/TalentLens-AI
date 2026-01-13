"""
Resume Parser Service
=====================

Service for extracting text and structure from PDF and DOCX files.
"""

import io
import re
from typing import Dict, Any, List, Optional
from loguru import logger

try:
    import PyPDF2
    from pdfminer.high_level import extract_text as pdfminer_extract
except ImportError:
    PyPDF2 = None
    pdfminer_extract = None

try:
    from docx import Document
except ImportError:
    Document = None


class ResumeParserService:
    """
    Service for parsing resume documents and extracting structured information.
    """
    
    # Common section headers
    SECTION_PATTERNS = {
        "experience": r"(?i)(work\s*experience|professional\s*experience|employment\s*history|experience)",
        "education": r"(?i)(education|academic\s*background|qualifications)",
        "skills": r"(?i)(skills|technical\s*skills|competencies|expertise)",
        "summary": r"(?i)(summary|profile|objective|about\s*me)",
        "projects": r"(?i)(projects|personal\s*projects|portfolio)",
        "certifications": r"(?i)(certifications|certificates|licenses)",
    }
    
    async def parse(self, file_content: bytes, file_ext: str) -> Dict[str, Any]:
        """
        Parse a resume file and extract structured information.
        
        Args:
            file_content: Raw file bytes
            file_ext: File extension (.pdf, .docx)
            
        Returns:
            Dictionary with extracted information
        """
        # Extract raw text
        if file_ext.lower() == ".pdf":
            raw_text = await self._extract_pdf_text(file_content)
        elif file_ext.lower() in [".docx", ".doc"]:
            raw_text = await self._extract_docx_text(file_content)
        else:
            raise ValueError(f"Unsupported file type: {file_ext}")
        
        # Parse sections
        sections = self._parse_sections(raw_text)
        
        # Extract structured data
        education = self._extract_education(sections.get("education", raw_text))
        work_history = self._extract_work_history(sections.get("experience", raw_text))
        experience_years = self._calculate_experience_years(work_history)
        
        return {
            "raw_text": raw_text,
            "sections": sections,
            "education": education,
            "work_history": work_history,
            "experience_years": experience_years
        }
    
    async def _extract_pdf_text(self, content: bytes) -> str:
        """Extract text from PDF file."""
        text = ""
        
        # Try PyPDF2 first
        if PyPDF2:
            try:
                pdf_file = io.BytesIO(content)
                reader = PyPDF2.PdfReader(pdf_file)
                for page in reader.pages:
                    text += page.extract_text() + "\n"
            except Exception as e:
                logger.warning(f"PyPDF2 extraction failed: {e}")
        
        # Fallback to pdfminer if text is empty or too short
        if len(text.strip()) < 100 and pdfminer_extract:
            try:
                pdf_file = io.BytesIO(content)
                text = pdfminer_extract(pdf_file)
            except Exception as e:
                logger.warning(f"pdfminer extraction failed: {e}")
        
        return self._clean_text(text)
    
    async def _extract_docx_text(self, content: bytes) -> str:
        """Extract text from DOCX file."""
        if not Document:
            raise ImportError("python-docx is required for DOCX parsing")
        
        try:
            doc_file = io.BytesIO(content)
            doc = Document(doc_file)
            
            text = ""
            for para in doc.paragraphs:
                text += para.text + "\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return self._clean_text(text)
            
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            raise
    
    def _clean_text(self, text: str) -> str:
        """Clean extracted text."""
        # Remove extra whitespace
        text = re.sub(r'\s+', ' ', text)
        # Remove special characters but keep punctuation
        text = re.sub(r'[^\w\s.,;:!?@#$%&*()\-+=\[\]{}\'\"\/]', '', text)
        return text.strip()
    
    def _parse_sections(self, text: str) -> Dict[str, str]:
        """Parse resume into sections based on common headers."""
        sections = {}
        
        # Find all section matches
        all_matches = []
        for section_name, pattern in self.SECTION_PATTERNS.items():
            for match in re.finditer(pattern, text):
                all_matches.append((match.start(), match.end(), section_name))
        
        # Sort by position
        all_matches.sort(key=lambda x: x[0])
        
        # Extract section content
        for i, (start, end, section_name) in enumerate(all_matches):
            if i < len(all_matches) - 1:
                next_start = all_matches[i + 1][0]
                sections[section_name] = text[end:next_start].strip()
            else:
                sections[section_name] = text[end:].strip()
        
        return sections
    
    def _extract_education(self, text: str) -> List[Dict[str, Any]]:
        """Extract education entries from text."""
        education = []
        
        # Common degree patterns
        degree_patterns = [
            r"(?i)(bachelor|b\.?s\.?|b\.?a\.?|b\.?sc\.?)",
            r"(?i)(master|m\.?s\.?|m\.?a\.?|m\.?sc\.?|mba)",
            r"(?i)(ph\.?d\.?|doctorate|doctor)",
            r"(?i)(associate|a\.?s\.?|a\.?a\.?)",
            r"(?i)(diploma|certificate)"
        ]
        
        # Year pattern
        year_pattern = r"20\d{2}|19\d{2}"
        
        # Find degrees
        for pattern in degree_patterns:
            matches = re.finditer(pattern, text)
            for match in matches:
                # Get surrounding context
                start = max(0, match.start() - 50)
                end = min(len(text), match.end() + 100)
                context = text[start:end]
                
                # Find year
                year_match = re.search(year_pattern, context)
                
                education.append({
                    "degree": match.group(),
                    "context": context.strip(),
                    "graduation_year": int(year_match.group()) if year_match else None
                })
        
        return education
    
    def _extract_work_history(self, text: str) -> List[Dict[str, Any]]:
        """Extract work history from text."""
        work_history = []
        
        # Date pattern for work experience
        date_pattern = r"((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s*\d{4}|\d{4})\s*[-–]\s*((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s*\d{4}|\d{4}|Present|Current)"
        
        matches = list(re.finditer(date_pattern, text, re.IGNORECASE))
        
        for i, match in enumerate(matches):
            start_date = match.group(1)
            end_date = match.group(2)
            
            # Get context around the date
            start = max(0, match.start() - 100)
            end = min(len(text), match.end() + 200)
            context = text[start:end]
            
            work_history.append({
                "start_date": start_date,
                "end_date": end_date,
                "context": context.strip()
            })
        
        return work_history
    
    def _calculate_experience_years(self, work_history: List[Dict]) -> int:
        """Calculate total years of experience from work history."""
        if not work_history:
            return 0
        
        import datetime
        total_months = 0
        
        for entry in work_history:
            try:
                # Parse dates
                start = self._parse_date(entry.get("start_date", ""))
                end_str = entry.get("end_date", "")
                
                if end_str.lower() in ["present", "current"]:
                    end = datetime.datetime.now()
                else:
                    end = self._parse_date(end_str)
                
                if start and end:
                    months = (end.year - start.year) * 12 + (end.month - start.month)
                    total_months += max(0, months)
            except:
                continue
        
        return total_months // 12
    
    def _parse_date(self, date_str: str) -> Optional[Any]:
        """Parse date string to datetime."""
        import datetime
        
        if not date_str:
            return None
        
        # Try year only
        year_match = re.search(r'(\d{4})', date_str)
        if year_match:
            try:
                return datetime.datetime(int(year_match.group(1)), 6, 1)
            except:
                pass
        
        return None
